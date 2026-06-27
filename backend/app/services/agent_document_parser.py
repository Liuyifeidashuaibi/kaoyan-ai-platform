"""
Unstructured 文档解析模块 — 商业级多格式文档解析。

核心能力：
  1. 多格式支持：docx, pdf, txt, md, csv, html, pptx, eml
  2. 结构化输出：标题、段落、列表、表格、图片引用
  3. 表格提取：从 PDF/DOCX 中提取表格为 Markdown 格式
  4. 元数据提取：文档属性、页数、标题层级
  5. 编码检测：自动检测文件编码（chardet）

降级策略：
  - unstructured 不可用时，降级到 python-docx / PyMuPDF / pdfplumber
  - 编码检测失败时，降级到 UTF-8
"""

from __future__ import annotations

import logging
import pathlib
from typing import Any

from app.config import get_settings

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    商业级文档解析器。

    使用 Unstructured 作为主引擎，自动降级到备选方案。
    """

    def __init__(self) -> None:
        self.settings = get_settings()
        self._unstructured_available = self._check_unstructured()

    def _check_unstructured(self) -> bool:
        """检查 unstructured 是否可用。"""
        try:
            from unstructured.partition.auto import partition
            logger.info("Unstructured 文档解析器可用")
            return True
        except ImportError:
            logger.warning("Unstructured 未安装，将使用备选解析器")
            return False
        except Exception as e:
            logger.warning("Unstructured 初始化失败: %s", e)
            return False

    def parse(
        self,
        file_path: str | pathlib.Path,
        max_chars: int = 20000,
    ) -> dict[str, Any]:
        """
        解析文档文件，返回结构化文本内容。

        返回:
            {
                "filename": str,
                "content": str,         # Markdown 格式文本
                "char_count": int,
                "elements_count": int,  # 解析出的元素数量
                "tables_count": int,    # 表格数量
                "format_info": str,     # 格式描述
                "doc_props": dict,      # 文档属性
                "summary": str,
            }
        """
        resolved = pathlib.Path(file_path)
        if not resolved.is_file():
            return {"error": f"文件不存在: {file_path}"}

        ext = resolved.suffix.lower()

        # 优先使用 Unstructured
        if self._unstructured_available:
            try:
                return self._parse_with_unstructured(resolved, ext, max_chars)
            except Exception as exc:
                logger.warning("Unstructured 解析失败，降级: %s", exc)

        # 降级到备选解析器
        return self._parse_with_fallback(resolved, ext, max_chars)

    def _parse_with_unstructured(
        self,
        file_path: pathlib.Path,
        ext: str,
        max_chars: int,
    ) -> dict[str, Any]:
        """使用 Unstructured 解析文档。"""
        from unstructured.partition.auto import partition

        # 根据扩展名选择分区策略
        elements = partition(filename=str(file_path))

        parts: list[str] = []
        tables_count = 0
        heading_count = 0

        for elem in elements:
            text = str(elem).strip()
            if not text:
                continue

            # 检测元素类型
            cat = elem.category if hasattr(elem, "category") else "Text"

            if cat == "Title" or cat == "Header":
                parts.append(f"# {text}")
                heading_count += 1
            elif cat == "Table":
                parts.append(f"| {text} |")
                tables_count += 1
            elif cat == "ListItem":
                parts.append(f"- {text}")
            elif cat == "NarrativeText" or cat == "Text":
                parts.append(text)
            else:
                parts.append(text)

        text = "\n".join(parts)

        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[... 内容已截断 ...]"

        return {
            "filename": file_path.name,
            "content": text,
            "char_count": len(text),
            "elements_count": len(elements),
            "tables_count": tables_count,
            "headings_count": heading_count,
            "format_info": f"Unstructured 解析（{len(elements)} 个元素，{heading_count} 个标题，{tables_count} 个表格）",
            "doc_props": {},
            "summary": f"已解析 {file_path.name}（{len(text)} 字符，{len(elements)} 个元素）",
        }

    def _parse_with_fallback(
        self,
        file_path: pathlib.Path,
        ext: str,
        max_chars: int,
    ) -> dict[str, Any]:
        """备选解析器（python-docx / PyMuPDF / pdfplumber）。"""
        text = ""
        format_info: list[str] = []
        doc_props: dict[str, str] = {}

        try:
            if ext in (".txt", ".md"):
                raw = file_path.read_bytes()
                encoding = self._detect_encoding(raw)
                text = raw.decode(encoding, errors="replace")
                format_info.append("纯文本格式")

            elif ext == ".csv":
                raw = file_path.read_bytes()
                encoding = self._detect_encoding(raw)
                text = raw.decode(encoding, errors="replace")
                format_info.append(f"CSV 格式（约 {text.count(chr(10)) + 1} 行）")

            elif ext == ".docx":
                from docx import Document
                doc = Document(str(file_path))

                cp = doc.core_properties
                if cp.title:
                    doc_props["title"] = cp.title
                if cp.author:
                    doc_props["author"] = cp.author
                if cp.created:
                    doc_props["created"] = str(cp.created)

                parts: list[str] = []
                heading_count = 0
                for para in doc.paragraphs:
                    if not para.text.strip():
                        continue
                    style_name = (para.style.name or "").lower()
                    if "heading 1" in style_name or "title" in style_name:
                        parts.append(f"# {para.text.strip()}")
                        heading_count += 1
                    elif "heading 2" in style_name:
                        parts.append(f"## {para.text.strip()}")
                        heading_count += 1
                    elif "heading 3" in style_name:
                        parts.append(f"### {para.text.strip()}")
                        heading_count += 1
                    elif "list" in style_name:
                        parts.append(f"- {para.text.strip()}")
                    else:
                        parts.append(para.text.strip())

                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):
                            parts.append("| " + " | ".join(cells) + " |")

                text = "\n".join(parts)
                format_info.append(f"DOCX（{heading_count} 个标题，{len(doc.tables)} 个表格）")
                if doc_props:
                    format_info.append(f"文档属性: {doc_props}")

            elif ext == ".pdf":
                try:
                    import fitz
                    doc = fitz.open(str(file_path))
                    parts = [page.get_text() for page in doc]
                    page_count = doc.page_count
                    doc.close()
                    text = "\n".join(parts)
                    format_info.append(f"PDF（{page_count} 页）")
                except ImportError:
                    import pdfplumber
                    parts: list[str] = []
                    with pdfplumber.open(str(file_path)) as pdf:
                        page_count = len(pdf.pages)
                        for page in pdf.pages:
                            t = page.extract_text()
                            if t:
                                parts.append(t)
                            # 尝试提取表格
                            tables = page.extract_tables()
                            for table in tables:
                                for row in table:
                                    cells = [c or "" for c in row]
                                    parts.append("| " + " | ".join(cells) + " |")
                    text = "\n".join(parts)
                    format_info.append(f"PDF（{page_count} 页，含表格提取）")

            else:
                return {"error": f"不支持的格式: {ext}"}
        except Exception as exc:
            logger.error("文档解析失败: %s", exc, exc_info=True)
            return {"error": f"解析失败: {exc}"}

        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[... 内容已截断 ...]"

        return {
            "filename": file_path.name,
            "content": text,
            "char_count": len(text),
            "elements_count": text.count("\n") + 1,
            "tables_count": text.count("| "),
            "format_info": "; ".join(format_info) if format_info else "未知",
            "doc_props": doc_props,
            "summary": f"已解析 {file_path.name}（{len(text)} 字符）",
        }

    def _detect_encoding(self, raw: bytes) -> str:
        """使用 chardet 自动检测文件编码。"""
        try:
            import chardet
            detected = chardet.detect(raw)
            encoding = detected.get("encoding", "utf-8") or "utf-8"
            # 修正常见误判
            if encoding.lower() in ("ascii", "iso-8859-1"):
                encoding = "utf-8"
            return encoding
        except ImportError:
            return "utf-8"
        except Exception:
            return "utf-8"


# 全局单例
_doc_parser: DocumentParser | None = None


def get_document_parser() -> DocumentParser:
    global _doc_parser
    if _doc_parser is None:
        _doc_parser = DocumentParser()
    return _doc_parser
