from pathlib import Path

from docx import Document as DocxDocument

from translator.core.exceptions import FileProcessingError
from translator.core.types import Document, InputKind
from translator.processors.base import InputProcessor
from translator.utils.source import extension_of

DOCX_EXTENSIONS = {".docx"}


class DocxProcessor:
    def can_handle(self, source: Path | str) -> bool:
        return extension_of(source) in DOCX_EXTENSIONS

    def extract(self, source: Path | str) -> Document:
        path = Path(source).resolve()
        if not path.is_file():
            raise FileNotFoundError(f"DOCX not found: {path}")

        try:
            doc = DocxDocument(path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        except Exception as exc:
            raise FileProcessingError(f"Failed to read DOCX: {path.name}") from exc

        if not paragraphs:
            raise FileProcessingError(f"No extractable text in DOCX: {path.name}")

        return Document(
            text="\n\n".join(paragraphs),
            source_name=path.name,
            needs_vision=False,
            kind=InputKind.DOCUMENT,
        )


def get_docx_processor() -> InputProcessor:
    return DocxProcessor()
